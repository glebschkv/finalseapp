"""Generate report_section_3.docx with full Section 3: Use Instructions."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# -- Page margins: 2cm all sides --
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

# -- Style defaults --
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)
paragraph_format = style.paragraph_format
paragraph_format.space_after = Pt(4)
paragraph_format.space_before = Pt(0)
paragraph_format.line_spacing = 1.0


def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color_hex,
    })
    shading.append(shading_elem)


def add_heading_styled(text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = "Calibri"
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(4)
    return h


def add_para(text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.size = Pt(11)
        run_b.font.name = "Calibri"
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    return p


def add_code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    return p


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Calibri"


def style_header_row(table):
    for cell in table.rows[0].cells:
        set_cell_shading(cell, "2E4057")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.bold = True


# ============================================================
# SECTION 3
# ============================================================
add_heading_styled("3 – Use Instructions", level=2)

add_para(
    "The source code for OBD InsightBot is available on GitHub at "
    "https://github.com/COMP2281/software-engineering-group25-26-18.git. "
    "A copy of the code has also been submitted via Ultra alongside this report."
)

# -- 3.1 Installation --
add_heading_styled("3.1 – Installation", level=3)

add_para(
    "Table 3.1 lists the minimum and recommended hardware specifications. "
    "A microphone is optional and only required for voice features (BR6, BR7)."
)

# System Requirements Table
req_table = doc.add_table(rows=5, cols=3)
req_table.style = "Table Grid"
headers = ["Component", "Minimum", "Recommended"]
data = [
    ["CPU", "4-core processor", "8-core processor"],
    ["RAM", "4 GB", "8 GB"],
    ["Storage", "4 GB free disk space", "8 GB free disk space"],
    ["GPU", "Not required", "CUDA 12+ compatible GPU"],
]
for i, h in enumerate(headers):
    req_table.rows[0].cells[i].text = h
for r, row_data in enumerate(data):
    for c, val in enumerate(row_data):
        req_table.rows[r + 1].cells[c].text = val
style_table(req_table)
style_header_row(req_table)

p_caption = doc.add_paragraph()
run_cap = p_caption.add_run("Table 3.1. Hardware requirements for OBD InsightBot")
run_cap.italic = True
run_cap.font.size = Pt(9)
run_cap.font.name = "Calibri"

# Supported OS Table
add_para("The application has been tested on the following operating systems:")

os_table = doc.add_table(rows=4, cols=2)
os_table.style = "Table Grid"
os_headers = ["Operating System", "Version"]
os_data = [
    ["Windows", "10 and 11"],
    ["macOS", "12 (Monterey) and above"],
    ["Linux", "Ubuntu 20.04+ / Debian-based distributions"],
]
for i, h in enumerate(os_headers):
    os_table.rows[0].cells[i].text = h
for r, row_data in enumerate(os_data):
    for c, val in enumerate(row_data):
        os_table.rows[r + 1].cells[c].text = val
style_table(os_table)
style_header_row(os_table)

p_cap2 = doc.add_paragraph()
run_cap2 = p_cap2.add_run("Table 3.2. Supported operating systems")
run_cap2.italic = True
run_cap2.font.size = Pt(9)
run_cap2.font.name = "Calibri"

add_para("Two prerequisites must be installed before setting up the application:")

p_prereq1 = doc.add_paragraph(style="List Bullet")
run1 = p_prereq1.add_run("Python 3.8 or higher")
run1.bold = True
r1b = p_prereq1.add_run(" – download from https://python.org. On Windows, tick "
                          '"Add Python to PATH" during installation.')
run1.font.size = Pt(11)
r1b.font.size = Pt(11)

p_prereq2 = doc.add_paragraph(style="List Bullet")
run2 = p_prereq2.add_run("Ollama")
run2.bold = True
r2b = p_prereq2.add_run(" – download from https://ollama.com. After installation, "
                          "pull the IBM Granite model:")
run2.font.size = Pt(11)
r2b.font.size = Pt(11)

add_code("ollama pull granite3.3:2b")

add_para(
    "Next, extract the submitted zip (or clone the repository) and run the "
    "following commands in a terminal:"
)

add_code("cd InsightBot")
add_code("python -m venv .venv                          # create virtual environment")

add_para("Activate the virtual environment:")

p_act = doc.add_paragraph(style="List Bullet")
p_act.add_run("Windows: ").bold = True
p_act.add_run(".venv\\Scripts\\activate").font.name = "Consolas"

p_act2 = doc.add_paragraph(style="List Bullet")
p_act2.add_run("macOS / Linux: ").bold = True
p_act2.add_run("source .venv/bin/activate").font.name = "Consolas"

add_para("Then install all dependencies:")
add_code("pip install -r requirements.txt")

# -- 3.2 Deployment --
add_heading_styled("3.2 – Deployment", level=3)

add_para(
    "OBD InsightBot is designed as a local-first desktop application. All processing — "
    "including AI inference, speech recognition, and data storage — runs entirely on "
    "the user's machine. No cloud services, containers, or virtual machines are required."
)

add_para(
    "Ollama must be running before the application is launched. On most systems, Ollama "
    "starts automatically as a background service after installation. If it does not, "
    "start it manually with:",
)
add_code("ollama serve")

add_para(
    "The SQLite database is created automatically at ",
    bold_prefix="Database setup. ",
)
p = doc.paragraphs[-1]
r = p.add_run("./data/obd_insightbot.db")
r.font.name = "Consolas"
r.font.size = Pt(10)
p.add_run(
    " on first launch. The directory is also created automatically. No manual "
    "database configuration is needed."
).font.size = Pt(11)

add_para(
    "For advanced users, the following environment variables can be set in a ",
    bold_prefix="Configuration. ",
)
p = doc.paragraphs[-1]
r = p.add_run(".env")
r.font.name = "Consolas"
r.font.size = Pt(10)
p.add_run(" file in the project root:").font.size = Pt(11)

env_table = doc.add_table(rows=5, cols=3)
env_table.style = "Table Grid"
env_headers = ["Variable", "Default", "Description"]
env_data = [
    ["OLLAMA_URL", "http://localhost:11434", "Ollama API endpoint"],
    ["OLLAMA_MODEL", "granite3.3:2b", "Language model name"],
    ["DATABASE_PATH", "./data/obd_insightbot.db", "SQLite database path"],
    ["APP_LOG_LEVEL", "INFO", "Logging verbosity (DEBUG, INFO, WARNING, ERROR)"],
]
for i, h in enumerate(env_headers):
    env_table.rows[0].cells[i].text = h
for r_idx, row_data in enumerate(env_data):
    for c, val in enumerate(row_data):
        cell = env_table.rows[r_idx + 1].cells[c]
        cell.text = val
        if c < 2:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Consolas"
                    run.font.size = Pt(9)
style_table(env_table)
style_header_row(env_table)

p_cap3 = doc.add_paragraph()
run_cap3 = p_cap3.add_run("Table 3.3. Optional environment variables")
run_cap3.italic = True
run_cap3.font.size = Pt(9)
run_cap3.font.name = "Calibri"

# -- 3.3 Launching --
add_heading_styled("3.3 – Launching", level=3)

add_para("With the virtual environment activated and Ollama running, launch the application:")
add_code("python src/main.py")

add_para(
    "On first launch, click "
    '"Register" to create an account. Usernames must be 3–50 characters and passwords '
    "must be at least 6 characters. Passwords are hashed with bcrypt before storage. "
    "Login attempts are rate-limited to 5 per 5-minute window to guard against brute-force attacks.",
    bold_prefix="Account creation. ",
)

add_para(
    "The application is a single-user local system with no distinct admin or role-based access. "
    "Each registered account's chats are private and protected by a session token. "
    "Multiple accounts can be created, and any account can be deleted from within the application, "
    "which cascade-deletes all associated chats and messages.",
    bold_prefix="User roles. ",
)

add_para(
    'After logging in, click "New Chat" and upload an OBD-II CSV file. '
    "Three sample log files are included in the repository root "
    "(demo_log.csv, demo_log_2.csv, demo_log_3.csv) "
    'for immediate testing. Useful first queries include "What\'s wrong with my vehicle?" '
    'or "Explain fault code P0300".',
    bold_prefix="First-time setup. ",
)

# -- 3.4 Troubleshooting --
add_heading_styled("3.4 – Troubleshooting", level=3)

add_para("Table 3.4 lists common issues encountered during installation and use.")

# Troubleshooting Table
ts_table = doc.add_table(rows=7, cols=3)
ts_table.style = "Table Grid"
ts_headers = ["Error / Symptom", "Cause", "Solution"]
ts_data = [
    ['"Ollama not detected"', "Ollama service is not running",
     "Run ollama serve in a separate terminal window"],
    ['"Model not found"', "Granite model has not been pulled",
     "Run ollama pull granite3.3:2b"],
    ['"python is not recognized"', "Python is not on the system PATH",
     "Reinstall Python and enable 'Add to PATH' during setup"],
    ['"No module named PyQt6"', "Dependencies not installed or venv not active",
     "Activate the virtual environment, then run pip install -r requirements.txt"],
    ["Microphone warning on startup", "No microphone connected or OS permissions denied",
     "Connect a microphone; on macOS/Linux, grant microphone permissions in system settings"],
    ["Status bar shows 'AI: Demo Mode'", "Ollama became unavailable at runtime",
     "Restart Ollama; the application will reconnect automatically"],
]
for i, h in enumerate(ts_headers):
    ts_table.rows[0].cells[i].text = h
for r_idx, row_data in enumerate(ts_data):
    for c, val in enumerate(row_data):
        ts_table.rows[r_idx + 1].cells[c].text = val
style_table(ts_table)
style_header_row(ts_table)

p_cap4 = doc.add_paragraph()
run_cap4 = p_cap4.add_run("Table 3.4. Common errors and their solutions")
run_cap4.italic = True
run_cap4.font.size = Pt(9)
run_cap4.font.name = "Calibri"

add_para(
    "The application writes daily log files to the ",
    bold_prefix="Logs. ",
)
p = doc.paragraphs[-1]
r = p.add_run("./logs/")
r.font.name = "Consolas"
r.font.size = Pt(10)
p.add_run(" directory:").font.size = Pt(11)

p_log1 = doc.add_paragraph(style="List Bullet")
r_l1 = p_log1.add_run("obd_insightbot_YYYYMMDD.log")
r_l1.font.name = "Consolas"
r_l1.font.size = Pt(10)
p_log1.add_run(" — full activity log (all severity levels)").font.size = Pt(11)

p_log2 = doc.add_paragraph(style="List Bullet")
r_l2 = p_log2.add_run("obd_insightbot_errors_YYYYMMDD.log")
r_l2.font.name = "Consolas"
r_l2.font.size = Pt(10)
p_log2.add_run(" — errors and critical issues only").font.size = Pt(11)

add_para(
    "Additionally, the application runs an internal health-check system that monitors "
    "database connectivity, AI backend availability, and disk space. Component status "
    "is shown via status indicators in the application interface, using three levels: "
    "Healthy (green), Degraded (amber), and Unhealthy (red)."
)

# -- Save --
doc.save("/home/user/finalseapp/report_section_3.docx")
print("report_section_3.docx generated successfully.")
